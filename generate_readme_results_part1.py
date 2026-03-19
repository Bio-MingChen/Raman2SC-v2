from pathlib import Path
from collections import defaultdict
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

base = Path(r"h:\coding_life\Academic_career\Raman_single_cell")
results_dir = base / "results"

nb_files = [
    base / "raman_pclda.ipynb",
    base / "Raman_B_basic_spectra.ipynb",
    base / "ABC_B_lineage_single_cell_analysis.ipynb",
]

FIG_DESC = {
    "raman_pclda.ipynb": {
        "pclda": "PC-LDA散点图（训练/测试或细胞类型可视化）",
        "lda1_lda2": "LDA1 vs LDA2散点图",
        "lda1_lda3": "LDA1 vs LDA3散点图",
        "lda_3d": "LDA 3D散点图",
        "umap": "UMAP可视化",
        "roc": "ROC曲线（分类性能）",
        "confusion": "混淆矩阵/分类统计表",
    },
    "Raman_B_basic_spectra.ipynb": {
        "raman_heatmap": "600–1800 cm^-1 拉曼光谱热图（分箱与归一化）",
        "raman_raw_spectra": "原始光谱分布（均值±标准差，高亮区间）",
        "raman_scaled_spectra": "标准化光谱分布（均值±标准差，高亮区间）",
    },
    "ABC_B_lineage_single_cell_analysis.ipynb": {
        "heatmap": "差异基因热图（细胞类型）",
        "tracksplot": "差异基因轨迹图（tracksplot）",
        "umap": "UMAP细胞类型可视化",
        "dotplot": "Marker基因dotplot",
        "kegg": "KEGG富集可视化",
        "go": "GO BP富集可视化",
        "lipid": "脂质相关富集可视化",
    },
}

summaries = {
    "raman_pclda.ipynb": "Raman 单细胞光谱的PC-LDA分析与可视化（含LDA散点/3D/UMAP、ROC与分类统计）。",
    "Raman_B_basic_spectra.ipynb": "Raman光谱基础比较（热图、原始与标准化光谱分布、差异/峰值汇总与表格输出）。",
    "ABC_B_lineage_single_cell_analysis.ipynb": "单细胞转录组分析（差异基因、UMAP、marker基因dotplot、KEGG/GO富集与脂质相关分析）。",
}

heading_titles = {
    "raman_pclda.ipynb": "Raman LDA analysis",
    "Raman_B_basic_spectra.ipynb": "Raman spectra analysis",
    "ABC_B_lineage_single_cell_analysis.ipynb": "Single cell analysis",
}

notebook_dirs = {
    "raman_pclda.ipynb": {"figure1", "figureS2"},
    "Raman_B_basic_spectra.ipynb": {"figure2", "figureS3"},
    "ABC_B_lineage_single_cell_analysis.ipynb": {"figure3"},
}

fig_files = []
for p in results_dir.rglob("*"):
    if p.is_file() and p.suffix.lower() in {".png", ".svg", ".pdf", ".docx", ".xlsx"}:
        fig_files.append(p)

by_nb = defaultdict(list)
for f in fig_files:
    rel = f.relative_to(base)
    parts = set(rel.parts)
    for nb, dirs in notebook_dirs.items():
        if parts & dirs:
            by_nb[nb].append(rel)


def describe_file(nb, relpath):
    name = relpath.stem.lower()
    desc_map = FIG_DESC.get(nb, {})
    for key, desc in desc_map.items():
        if key in name:
            return desc
    if relpath.suffix.lower() == ".xlsx":
        return "结果表格/统计输出"
    if relpath.suffix.lower() == ".docx":
        return "结果文档输出"
    return "结果图或输出文件"

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    new_run.append(r_pr)
    text_elm = OxmlElement("w:t")
    text_elm.text = text
    new_run.append(text_elm)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

out_path = results_dir / "README_results_part1.docx"

doc = Document()
doc.add_heading("Analysis Summary (README_results_part1)", level=1)
doc.add_paragraph("")

for nb in nb_files:
    nb_name = nb.name
    title = heading_titles.get(nb_name, nb_name)
    doc.add_heading(title, level=2)
    doc.add_paragraph(summaries.get(nb_name, "分析结果摘要。"))
    items = sorted(by_nb.get(nb_name, []), key=lambda x: str(x))
    if items:
        doc.add_paragraph("Figures/Outputs:")
        for rel in items:
            desc = describe_file(nb_name, rel)
            p = doc.add_paragraph(style="List Bullet")
            add_hyperlink(p, str((base / rel).resolve()), rel.as_posix())
            p.add_run(f": {desc}")
    else:
        doc.add_paragraph("Figures/Outputs: (none detected in results)")
    doc.add_paragraph("")

doc.save(out_path)
print(f"Saved: {out_path}")
